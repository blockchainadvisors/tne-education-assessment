"""
Generate the TNE Assessment Platform - Project Scoping & Requirements Document (.docx)
Run: cd backend && ../.venv/bin/python -m scripts.generate_scoping_doc
  OR: cd /home/neylaur/tne-education-assessment/backend && .venv/bin/python ../scripts/generate_scoping_doc.py
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
from datetime import date


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a consistently styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "4F46E5")

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F1F5F9")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def add_heading_with_number(doc, text, level=1):
    """Add a numbered heading."""
    doc.add_heading(text, level=level)


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(text)
        run2.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def build_document():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Heading styles
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        hs.font.name = "Calibri"

    # ═══════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TNE Quality Assessment &\nBenchmarking Platform")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Project Scoping & Requirements Document")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Version 1.0  |  {date.today().strftime('%B %d, %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph("")
    doc.add_paragraph("")

    status_line = doc.add_paragraph()
    status_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = status_line.add_run("CONFIDENTIAL")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (placeholder)
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1.  Executive Summary",
        "2.  Project Overview",
        "3.  Functional Requirements",
        "    3.1  Authentication & User Management",
        "    3.2  Multi-Tenancy & Organisation Management",
        "    3.3  Assessment Workflow",
        "    3.4  AI-Powered Capabilities",
        "    3.5  Benchmarking & Analytics",
        "    3.6  File Management & Document Intelligence",
        "    3.7  Reporting",
        "    3.8  Administration",
        "4.  Data Model",
        "5.  API Specification",
        "6.  Frontend Application",
        "    6.1  Authentication Pages",
        "    6.2  Dashboard",
        "    6.3  Assessment Form",
        "    6.4  Review & Scores",
        "    6.5  Benchmarks",
        "    6.6  Administration",
        "7.  AI & Machine Learning Pipeline",
        "    7.1  Text Scoring Engine",
        "    7.2  Numeric & Binary Scoring",
        "    7.3  Timeseries Analysis",
        "    7.4  Consistency Checking",
        "    7.5  Report Generation",
        "    7.6  Document Intelligence",
        "    7.7  Risk Prediction",
        "8.  Infrastructure & Deployment",
        "9.  Testing Strategy",
        "10. Non-Functional Requirements",
        "11. Assessment Template Specification",
        "12. Glossary",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("1. Executive Summary", level=1)

    add_body(doc,
        "The TNE Quality Assessment & Benchmarking Platform is a multi-tenant SaaS application "
        "designed to streamline the evaluation and quality assurance of transnational education (TNE) "
        "partnerships. The platform enables institutions to complete structured self-assessments across "
        "52 items spanning 5 thematic areas, leverages AI to automatically score responses and generate "
        "comprehensive quality reports, and provides benchmarking capabilities to compare performance "
        "against peer institutions globally and regionally."
    )

    add_body(doc,
        "The system is built on a modern technology stack comprising a FastAPI backend with PostgreSQL, "
        "a Next.js 15 frontend with React 19, AI-powered analysis via large language models, and "
        "asynchronous task processing via Celery and Redis. The platform supports five user roles, "
        "complete assessment lifecycle management (draft through to report generation), and automated "
        "document intelligence for uploaded evidence files."
    )

    doc.add_heading("Key Deliverables", level=2)
    deliverables = [
        ("Multi-tenant platform: ", "Secure, isolated environments for each institution with role-based access control."),
        ("52-item assessment framework: ", "Structured across 5 weighted themes covering Teaching & Learning, Student Experience, Governance, Impact, and Finance."),
        ("AI-powered scoring: ", "Automatic evaluation of text responses against quality rubrics across 4 dimensions, plus rule-based scoring for numeric, binary, and timeseries data."),
        ("Automated report generation: ", "Executive summaries, per-theme deep analysis, and prioritised improvement recommendations."),
        ("Document intelligence: ", "Automatic classification, completeness checking, and data extraction from uploaded PDF/DOCX files."),
        ("Peer benchmarking: ", "Percentile-based comparison against institutions by country and globally."),
        ("Risk prediction: ", "Rule-based early warning system identifying partnerships at risk across 6 weighted factors."),
    ]
    for bold, rest in deliverables:
        add_bullet(doc, rest, bold_prefix=bold)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 2. PROJECT OVERVIEW
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("2. Project Overview", level=1)

    doc.add_heading("2.1 Problem Statement", level=2)
    add_body(doc,
        "Transnational education partnerships require rigorous quality assurance processes that are "
        "currently manual, inconsistent, and time-consuming. Institutions lack standardised frameworks "
        "for self-assessment, peer comparison, and evidence-based reporting. The absence of automated "
        "analysis means quality reviews are subjective and resource-intensive."
    )

    doc.add_heading("2.2 Solution", level=2)
    add_body(doc,
        "A cloud-based platform that digitises the entire assessment lifecycle \u2014 from data collection "
        "through AI-powered analysis to report generation and benchmarking \u2014 providing institutions "
        "with objective, consistent, and actionable quality insights."
    )

    doc.add_heading("2.3 Technology Stack", level=2)
    add_styled_table(doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "Next.js 15 + React 19 + TypeScript", "Single-page application with App Router, server-side rendering"],
            ["Styling", "Tailwind CSS + Recharts", "Utility-first CSS framework with data visualisation charts"],
            ["Backend API", "FastAPI + Python 3.11", "High-performance async REST API"],
            ["Database", "PostgreSQL 16 + SQLAlchemy (async)", "Relational database with JSONB support for flexible data"],
            ["ORM / Migrations", "SQLAlchemy 2.0 + Alembic", "Async ORM with schema version control"],
            ["Cache / Queue", "Redis 7", "Session caching, Celery task broker, rate limiting"],
            ["Task Processing", "Celery", "Asynchronous background job execution"],
            ["Object Storage", "MinIO (S3-compatible)", "Document and file uploads"],
            ["AI Engine", "LLM API (Anthropic SDK)", "Text scoring, report generation, document intelligence"],
            ["Authentication", "JWT (HS256)", "Stateless token-based auth with refresh tokens"],
            ["Email", "SMTP (Mailpit for dev)", "Verification emails, magic links, notifications"],
            ["Containerisation", "Docker + Docker Compose", "Multi-service orchestration for development and production"],
            ["CI/CD", "GitHub Actions", "Automated linting, testing, and build verification"],
            ["E2E Testing", "Playwright", "Browser-based end-to-end test automation"],
        ],
        col_widths=[1.5, 2.8, 2.2]
    )

    doc.add_heading("2.4 User Roles", level=2)
    add_styled_table(doc,
        ["Role", "Description", "Key Permissions"],
        [
            ["Platform Admin", "System-wide administrator", "Manage all tenants, view platform statistics, full access"],
            ["Tenant Admin", "Organisation administrator", "Manage users, partners, tenant settings, trigger AI jobs"],
            ["Assessor", "Assessment author", "Create and fill assessments, upload documents"],
            ["Reviewer", "Quality reviewer", "Review scored assessments, trigger reports, view benchmarks"],
            ["Institution User", "Read-only stakeholder", "View assessments, reports, and benchmarks"],
        ],
        col_widths=[1.3, 2.2, 3.0]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 3. FUNCTIONAL REQUIREMENTS
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("3. Functional Requirements", level=1)

    # 3.1 Auth
    doc.add_heading("3.1 Authentication & User Management", level=2)
    add_body(doc, "The platform implements a comprehensive authentication system with multiple sign-in methods and email verification.")

    features_auth = [
        ("Email/password registration: ", "New users register with email, password (min 8 characters), full name, and organisation details. A new tenant is created automatically."),
        ("Email verification: ", "Registration triggers a verification email with a 24-hour single-use token. Users cannot log in until verified."),
        ("Password login: ", "Standard email + password authentication returning JWT access (15-min) and refresh (7-day) tokens."),
        ("Magic link login: ", "Passwordless authentication via email link with 15-minute expiry, rate-limited to 3 requests per 60 seconds."),
        ("Token refresh: ", "Transparent token refresh on 401 responses with mutex pattern to prevent concurrent refresh races."),
        ("Rate limiting: ", "Redis-based rate limiting on verification emails and magic links (3 per 60 seconds per email)."),
        ("User management: ", "Tenant admins can create, update, and deactivate users within their organisation."),
    ]
    for bold, rest in features_auth:
        add_bullet(doc, rest, bold_prefix=bold)

    # 3.2 Multi-tenancy
    doc.add_heading("3.2 Multi-Tenancy & Organisation Management", level=2)
    features_mt = [
        ("Tenant isolation: ", "All data tables include tenant_id. Application-layer filtering ensures strict data isolation between organisations."),
        ("Partner institutions: ", "Each tenant can register up to 5 partner institutions with name, country, and display ordering."),
        ("Tenant settings: ", "JSONB-based flexible settings per tenant. Subscription tiers: free, basic, professional, enterprise."),
        ("Organisation profile: ", "Tenant admins manage institution name, country, and configuration settings."),
    ]
    for bold, rest in features_mt:
        add_bullet(doc, rest, bold_prefix=bold)

    # 3.3 Assessment
    doc.add_heading("3.3 Assessment Workflow", level=2)
    add_body(doc, "Assessments follow a defined lifecycle with clear status transitions:")

    add_styled_table(doc,
        ["Status", "Description", "Allowed Transitions"],
        [
            ["Draft", "Institution is filling in responses", "Submitted"],
            ["Submitted", "Awaiting AI processing or reviewer action", "Under Review"],
            ["Under Review", "Being reviewed by a reviewer", "Scored"],
            ["Scored", "All responses have been AI-scored", "Report Generated"],
            ["Report Generated", "Full report is available (terminal state)", "\u2014"],
        ],
        col_widths=[1.5, 3.0, 2.0]
    )

    doc.add_paragraph("")
    features_assess = [
        ("Template-based: ", "Assessments are created from versioned templates containing 52 items across 5 weighted themes."),
        ("12 field types: ", "short_text, long_text, numeric, percentage, yes_no_conditional, dropdown, multi_select, file_upload, multi_year_gender, partner_specific, auto_calculated, salary_bands."),
        ("Auto-save: ", "Responses auto-save with 2-second debounce. Bulk save endpoint for batch operations."),
        ("Auto-calculations: ", "Derived metrics (student-staff ratio, PhD%, retention rate, etc.) computed automatically from dependent fields."),
        ("Submission validation: ", "All required items must have responses before submission is allowed."),
        ("One assessment per year: ", "Unique constraint on (tenant_id, academic_year) prevents duplicate assessments."),
    ]
    for bold, rest in features_assess:
        add_bullet(doc, rest, bold_prefix=bold)

    # 3.4 AI
    doc.add_heading("3.4 AI-Powered Capabilities", level=2)
    add_body(doc, "Four distinct AI capabilities are integrated into the platform, all processed asynchronously via background workers:")

    ai_caps = [
        ("Assessment scoring: ", "Evaluates all 52 items using type-specific scorers. Text responses are scored against quality rubrics across 4 dimensions (relevance, specificity, evidence, comprehensiveness). Numeric, binary, and timeseries data use rule-based scoring."),
        ("Report generation: ", "Three-stage AI pipeline producing an executive summary (~500 words), per-theme deep analysis (~300 words each), and 6\u20138 prioritised improvement recommendations with timelines."),
        ("Document intelligence: ", "Uploaded documents are automatically classified into 10 categories, checked for completeness against assessment requirements, and key data is extracted. Supports PDF and DOCX formats."),
        ("Risk prediction: ", "Identifies at-risk partnerships using 6 weighted factors: financial health (25%), enrollment trends (20%), student retention (15%), staff-student ratios (15%), governance strength (15%), and staff qualifications (10%)."),
    ]
    for bold, rest in ai_caps:
        add_bullet(doc, rest, bold_prefix=bold)

    # 3.5 Benchmarking
    doc.add_heading("3.5 Benchmarking & Analytics", level=2)
    features_bench = [
        ("Percentile comparison: ", "Institution scores compared against 10th, 25th, 50th, 75th, and 90th percentile benchmarks."),
        ("Geographic filtering: ", "Compare against institutions in the same country or globally."),
        ("Per-theme breakdown: ", "Benchmarks available for each of the 5 assessment themes individually."),
        ("Sample size transparency: ", "Each benchmark metric shows the number of institutions in the comparison set."),
        ("Visualisation: ", "Line charts comparing institution score against peer median with percentile bands."),
    ]
    for bold, rest in features_bench:
        add_bullet(doc, rest, bold_prefix=bold)

    # 3.6 Files
    doc.add_heading("3.6 File Management & Document Intelligence", level=2)
    features_files = [
        ("Upload: ", "Multipart file upload to S3-compatible storage (MinIO). Files stored with tenant/assessment/file path structure."),
        ("Download: ", "Presigned URLs with 1-hour TTL for secure, direct-from-storage downloads."),
        ("Automatic processing: ", "Upload triggers asynchronous document intelligence pipeline via Celery task."),
        ("Classification: ", "AI classifies documents into 10 categories (policy, financial report, meeting minutes, etc.) with confidence scores. Filename-based fallback when AI is unavailable."),
        ("Completeness checking: ", "AI evaluates whether uploaded documents satisfy assessment item requirements, scoring 0\u2013100 with specific section-level feedback."),
    ]
    for bold, rest in features_files:
        add_bullet(doc, rest, bold_prefix=bold)

    # 3.7 Reporting
    doc.add_heading("3.7 Reporting", level=2)
    features_report = [
        ("Version tracking: ", "Each report generation creates a new version, preserving history."),
        ("Executive summary: ", "AI-generated ~500-word overview covering performance context, key strengths, areas for improvement, and forward-looking conclusions."),
        ("Theme analysis: ", "~300-word deep dive for each of the 5 themes, covering strongest/weakest items and specific recommendations."),
        ("Improvement recommendations: ", "6\u20138 prioritised recommendations with clear titles, priority levels (High/Medium/Low), affected themes, data-backed rationale, and suggested timelines."),
        ("PDF export: ", "Planned HTML-to-PDF rendering via WeasyPrint with S3 storage."),
        ("Report viewer: ", "Collapsible section-based viewer in the frontend with expand/collapse all functionality."),
    ]
    for bold, rest in features_report:
        add_bullet(doc, rest, bold_prefix=bold)

    # 3.8 Admin
    doc.add_heading("3.8 Administration", level=2)
    features_admin = [
        ("Platform statistics: ", "Total tenants, users, and assessments (platform_admin only)."),
        ("Tenant listing: ", "View all tenants with subscription tiers and activity status."),
        ("User management: ", "Create, update, and deactivate users within a tenant."),
        ("Partner management: ", "CRUD operations for partner institutions (max 5 per tenant)."),
    ]
    for bold, rest in features_admin:
        add_bullet(doc, rest, bold_prefix=bold)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 4. DATA MODEL
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("4. Data Model", level=1)
    add_body(doc, "The platform uses 13 PostgreSQL tables with UUID primary keys and JSONB for flexible data storage. All data tables enforce multi-tenant isolation via tenant_id foreign keys.")

    doc.add_heading("4.1 Entity Summary", level=2)
    add_styled_table(doc,
        ["Entity", "Table", "Description", "Key Fields"],
        [
            ["User", "users", "Platform users with role-based access", "email, role, tenant_id, email_verified"],
            ["Tenant", "tenants", "Organisations/institutions", "name, slug, country, subscription_tier, settings (JSONB)"],
            ["Partner Institution", "partner_institutions", "TNE partner institutions per tenant", "name, country, position"],
            ["Assessment Template", "assessment_templates", "Versioned assessment structure", "name, version, is_active"],
            ["Assessment Theme", "assessment_themes", "5 thematic areas within template", "name, slug, weight, display_order"],
            ["Assessment Item", "assessment_items", "52 individual questions/fields", "code, field_type, field_config (JSONB), scoring_rubric (JSONB)"],
            ["Assessment", "assessments", "Institution assessment instance", "academic_year, status, overall_score, tenant_id"],
            ["Assessment Response", "assessment_responses", "Individual item responses", "value (JSONB), ai_score, ai_feedback"],
            ["Theme Score", "theme_scores", "Aggregated theme-level scores", "normalised_score, weighted_score, ai_analysis"],
            ["Assessment Report", "assessment_reports", "AI-generated quality reports", "executive_summary, theme_analyses (JSONB), recommendations (JSONB)"],
            ["File Upload", "file_uploads", "Uploaded documents", "storage_key, document_type, extraction_status, extracted_data (JSONB)"],
            ["Benchmark Snapshot", "benchmark_snapshots", "Percentile benchmarks", "percentile_10/25/50/75/90, sample_size"],
            ["AI Job", "ai_jobs", "Background task tracking", "job_type, status, progress, result_data (JSONB)"],
        ],
        col_widths=[1.3, 1.6, 1.8, 1.8]
    )

    doc.add_heading("4.2 Key Relationships", level=2)
    relationships = [
        "Tenant \u2192 Users (1:many), Partner Institutions (1:many), Assessments (1:many)",
        "Assessment Template \u2192 Themes (1:many) \u2192 Items (1:many)",
        "Assessment \u2192 Responses (1:many), Theme Scores (1:many), Report (1:1), AI Jobs (1:many)",
        "Assessment Response \u2192 Item (many:1), optionally Partner Institution (many:1)",
        "Unique constraints: (tenant_id, academic_year) on assessments; (assessment_id, item_id, partner_id) on responses",
    ]
    for r in relationships:
        add_bullet(doc, r)

    doc.add_heading("4.3 Assessment Field Types", level=2)
    add_styled_table(doc,
        ["Field Type", "Description", "Scoring Method", "Example Items"],
        [
            ["short_text", "Single-line text input", "AI (4-dimension rubric)", "Programme names, descriptions"],
            ["long_text", "Multi-line narrative text", "AI (4-dimension rubric)", "Quality assurance processes, strategies"],
            ["numeric", "Integer/decimal number", "Rule-based (range mapping)", "Staff count, programme count"],
            ["percentage", "0\u2013100% value", "Rule-based (range mapping)", "Retention rate, employment rate"],
            ["yes_no_conditional", "Boolean with follow-up text", "Binary + evidence quality", "Policy existence with details"],
            ["dropdown", "Single selection", "Skipped", "Accreditation body selection"],
            ["multi_select", "Multiple selections", "Skipped", "Quality frameworks adopted"],
            ["file_upload", "Document upload (PDF/DOCX)", "Document intelligence", "Policy documents, agreements"],
            ["multi_year_gender", "Time series with gender breakdown", "Trend analysis", "4-year enrollment data by gender"],
            ["partner_specific", "Per-partner institution data", "Skipped", "Partner-level metrics"],
            ["auto_calculated", "Formula-derived field", "Rule-based (range mapping)", "Student-staff ratio, PhD%"],
            ["salary_bands", "Staff compensation table", "Skipped", "Professor/Associate salary ranges"],
        ],
        col_widths=[1.4, 1.8, 1.6, 1.7]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 5. API SPECIFICATION
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("5. API Specification", level=1)
    add_body(doc, "The backend exposes a RESTful API at /api/v1 with 11 routers and 38+ endpoints. All authenticated endpoints require a valid JWT Bearer token.")

    doc.add_heading("5.1 Endpoint Summary", level=2)
    api_rows = [
        ["POST", "/auth/login", "Public", "Authenticate with email/password"],
        ["POST", "/auth/register", "Public", "Register new tenant and first admin user"],
        ["POST", "/auth/refresh", "Public", "Exchange refresh token for new token pair"],
        ["POST", "/auth/verify-email", "Public", "Verify email address via token"],
        ["POST", "/auth/magic-link", "Public", "Request passwordless login link"],
        ["POST", "/auth/magic-link/verify", "Public", "Login via magic link token"],
        ["POST", "/auth/resend-verification", "Public", "Resend verification email (rate-limited)"],
        ["GET", "/users/me", "Any", "Get current user profile"],
        ["GET", "/users", "Admin", "List users in tenant"],
        ["POST", "/users", "Admin", "Create user in tenant"],
        ["PUT", "/users/{id}", "Admin", "Update user"],
        ["GET", "/tenants/current", "Any", "Get current tenant details"],
        ["PUT", "/tenants/current", "Admin", "Update tenant settings"],
        ["GET", "/tenants/current/partners", "Any", "List partner institutions"],
        ["POST", "/tenants/current/partners", "Admin", "Add partner institution"],
        ["PUT", "/tenants/current/partners/{id}", "Admin", "Update partner"],
        ["DELETE", "/tenants/current/partners/{id}", "Admin", "Soft-delete partner"],
        ["GET", "/assessments/templates", "Any", "List assessment templates"],
        ["GET", "/assessments/templates/{id}", "Any", "Get template with themes and items"],
        ["GET", "/assessments", "Any", "List tenant assessments"],
        ["POST", "/assessments", "Any", "Create new assessment"],
        ["GET", "/assessments/{id}", "Any", "Get assessment details"],
        ["POST", "/assessments/{id}/submit", "Any", "Submit assessment for review"],
        ["POST", "/assessments/{id}/status/{status}", "Admin/Reviewer", "Change assessment status"],
        ["GET", "/assessments/{id}/responses", "Any", "List all responses"],
        ["PUT", "/assessments/{id}/responses/{item_id}", "Any", "Save single response (auto-calc)"],
        ["PUT", "/assessments/{id}/responses", "Any", "Bulk save responses"],
        ["GET", "/assessments/{id}/scores", "Any", "Get assessment scores"],
        ["POST", "/assessments/{id}/scores/trigger-scoring", "Admin/Reviewer", "Start AI scoring (async)"],
        ["GET", "/assessments/{id}/report", "Any", "Get assessment report"],
        ["POST", "/assessments/{id}/report/generate", "Admin/Reviewer", "Generate AI report (async)"],
        ["POST", "/assessments/{id}/files", "Any", "Upload file"],
        ["GET", "/assessments/{id}/files", "Any", "List uploaded files"],
        ["GET", "/assessments/{id}/files/{file_id}", "Any", "Get file metadata"],
        ["GET", "/assessments/{id}/files/{file_id}/download", "Any", "Get presigned download URL"],
        ["GET", "/benchmarks/compare/{id}", "Any", "Compare against peer benchmarks"],
        ["GET", "/admin/stats", "Platform Admin", "Platform-wide statistics"],
        ["GET", "/admin/tenants", "Platform Admin", "List all tenants"],
        ["GET", "/jobs/{id}", "Any", "Poll async AI job status"],
    ]
    add_styled_table(doc,
        ["Method", "Path", "Auth", "Description"],
        api_rows,
        col_widths=[0.7, 3.0, 1.2, 1.6]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 6. FRONTEND APPLICATION
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("6. Frontend Application", level=1)
    add_body(doc, "The frontend is a Next.js 15 application using the App Router pattern with React 19 and TypeScript. It uses React Query for server state management and Tailwind CSS for styling.")

    doc.add_heading("6.1 Authentication Pages", level=2)
    auth_pages = [
        ["/login", "Email/password login with magic link tab", "Public"],
        ["/register", "Organisation registration with auto-slug generation", "Public"],
        ["/verify-email", "Token-based email verification with auto-login", "Public"],
        ["/verify-email-sent", "Confirmation page with resend option", "Public"],
        ["/magic-link", "Magic link token verification", "Public"],
    ]
    add_styled_table(doc,
        ["Route", "Description", "Access"],
        auth_pages,
        col_widths=[1.5, 3.5, 1.5]
    )

    doc.add_heading("6.2 Dashboard", level=2)
    add_body(doc, "The main dashboard provides a comprehensive overview of the institution's assessment status and performance:")
    dash_features = [
        ("Status cards: ", "4-column grid showing counts for Draft, Under Review, Scored, and Completed assessments."),
        ("Overall score gauge: ", "Custom SVG semi-circle gauge displaying the latest assessment score with performance label (Strong/Developing/Needs Improvement)."),
        ("Theme radial bars: ", "Small gauge charts for each of the 5 themes."),
        ("Year-over-year trend: ", "Area chart showing score progression across academic years with 70% target reference line."),
        ("Benchmark comparison: ", "Line chart comparing institution score against peer median."),
        ("Recent assessments table: ", "5 most recent assessments with status badges, scores, and action links."),
        ("Radar chart: ", "Collapsible radar visualisation of theme performance (shown when 3+ themes scored)."),
    ]
    for bold, rest in dash_features:
        add_bullet(doc, rest, bold_prefix=bold)

    doc.add_heading("6.3 Assessment Form", level=2)
    add_body(doc, "The assessment editing interface supports all 12 field types with specialised renderers:")
    form_features = [
        ("Theme navigation: ", "Desktop sidebar or mobile dropdown showing completion counts (X/Y items per theme)."),
        ("Progress tracking: ", "Overall completion percentage bar."),
        ("Auto-save: ", "2-second debounce saves responses automatically with visual status indicator (Idle/Saving/Saved/Error)."),
        ("Field renderers: ", "12 specialised components for each field type including drag-and-drop file upload, multi-year gender tables, partner-specific data grids, and salary band matrices."),
        ("Conditional logic: ", "Fields with depends_on are shown/hidden based on parent responses."),
        ("Submit confirmation: ", "Final theme shows submit button with confirmation dialog."),
    ]
    for bold, rest in form_features:
        add_bullet(doc, rest, bold_prefix=bold)

    doc.add_heading("6.4 Review & Scores", level=2)
    review_features = [
        ("Read-only display: ", "All responses rendered in human-readable format using ResponseDisplay component."),
        ("Score summary: ", "Overall percentage with per-theme breakdowns showing normalised and weighted scores."),
        ("AI feedback: ", "Item-level AI analysis and dimension scores visible for text responses."),
        ("Report viewer: ", "Collapsible section-based viewer with expand/collapse all for executive summary, theme analyses, and recommendations."),
    ]
    for bold, rest in review_features:
        add_bullet(doc, rest, bold_prefix=bold)

    doc.add_heading("6.5 Benchmarks", level=2)
    bench_features = [
        ("Assessment selector: ", "Dropdown filtered to scored/completed assessments."),
        ("Country filter: ", "Compare against same-country peers or global benchmarks."),
        ("Comparison chart: ", "Line chart with institution score vs. peer median for each metric."),
        ("Percentile tooltip: ", "Custom tooltip showing P25, P75 bands and sample sizes."),
    ]
    for bold, rest in bench_features:
        add_bullet(doc, rest, bold_prefix=bold)

    doc.add_heading("6.6 Administration", level=2)
    add_body(doc, "Admin pages are restricted to platform_admin and tenant_admin roles. Currently provides 4 feature areas: Institution Settings, User Management, Partner Institutions, and System Configuration.")

    doc.add_heading("6.7 UI Component Library", level=2)
    add_styled_table(doc,
        ["Component", "Purpose"],
        [
            ["GaugeChart", "Custom SVG semi-circle gauge with needle animation, ticks, and colour variants"],
            ["StatusBadge", "Assessment status pills with semantic colours"],
            ["Alert", "Error/success/warning/info message boxes"],
            ["Badge", "Inline labels in 7 colour variants"],
            ["PageHeader", "Page title + description + action buttons"],
            ["EmptyState", "Illustrated empty data state with CTA"],
            ["ConfirmDialog", "Modal confirmation with focus trap and keyboard support"],
            ["Spinner", "SVG-based loading indicator in 3 sizes"],
            ["Logo", "Brand logo in full/mark variants with colour/white schemes"],
        ],
        col_widths=[1.5, 5.0]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 7. AI & ML PIPELINE
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("7. AI & Machine Learning Pipeline", level=1)
    add_body(doc,
        "All AI operations are coordinated through a centralised client wrapper with 7-day SHA256 in-memory caching, "
        "3-attempt exponential backoff retry (2s\u201330s), and per-call cost tracking. AI tasks execute asynchronously "
        "via Celery workers, with job status trackable through a polling endpoint."
    )

    doc.add_heading("7.1 Text Scoring Engine", level=2)
    add_body(doc, "Text responses (short_text and long_text) are evaluated against a 4-dimension rubric, each scored 0\u201325 points for a total of 0\u2013100:")
    add_styled_table(doc,
        ["Dimension", "Weight", "Evaluates"],
        [
            ["Relevance", "0\u201325", "How directly the response addresses the question"],
            ["Specificity", "0\u201325", "Presence of concrete examples, data, and details"],
            ["Evidence", "0\u201325", "Supporting data, processes, and external validation"],
            ["Comprehensiveness", "0\u201325", "Complete coverage of all aspects of the question"],
        ],
        col_widths=[1.5, 1.0, 4.0]
    )
    doc.add_paragraph("")
    add_body(doc, "Parameters: temperature=0.0 (deterministic), max_tokens=1000, caching enabled. Minimum input length: 10 characters; text truncated to 3,000 characters for API calls.")

    doc.add_heading("7.2 Numeric & Binary Scoring", level=2)
    add_body(doc, "Rule-based scoring with no AI API calls:")
    num_features = [
        ("Numeric scorer: ", "Maps values to predefined score ranges from the item's rubric (e.g., 1\u20135 programmes = 40, 5\u201315 = 70, 15+ = 100)."),
        ("Binary scorer: ", "Yes = 100, No = 0. When evidence text is provided, combines 30% binary + 70% evidence quality (scored by text length: <200 chars = 40, 200\u2013500 = 65, 500+ = 85)."),
    ]
    for bold, rest in num_features:
        add_bullet(doc, rest, bold_prefix=bold)

    doc.add_heading("7.3 Timeseries Analysis", level=2)
    add_body(doc, "Multi-year gender data is analysed for trends using linear regression. Base score of 60 is adjusted by \u00b120 depending on whether the trend direction matches the rubric's ideal direction (e.g., increasing enrollment is positive). Supports both gendered (male/female) and simple numeric time series formats.")

    doc.add_heading("7.4 Consistency Checking", level=2)
    add_body(doc, "Two-phase validation of assessment data integrity:")
    consist_features = [
        ("Phase 1 \u2013 Rule-based: ", "Predefined rules check logical constraints (e.g., PhD staff \u2264 total staff, retention rate 0\u2013100%, flying faculty \u2264 total staff)."),
        ("Phase 2 \u2013 AI analysis: ", "Optional LLM-based cross-theme inconsistency detection reviewing the full assessment data (limited to 5,000 characters). Identifies contradictions, implausible claims, and missing evidence."),
    ]
    for bold, rest in consist_features:
        add_bullet(doc, rest, bold_prefix=bold)

    doc.add_heading("7.5 Report Generation", level=2)
    add_body(doc, "Three-stage AI pipeline with creative prose generation (temperature=0.3):")
    add_styled_table(doc,
        ["Stage", "Output", "Max Tokens", "Details"],
        [
            ["Executive Summary", "~500-word overview", "2,000", "Performance context, 2\u20133 strengths, 2\u20133 improvements, forward-looking conclusion"],
            ["Theme Analysis (\u00d75)", "~300 words per theme", "1,500 each", "Strongest/weakest items, benchmark comparison, specific recommendations"],
            ["Recommendations", "6\u20138 prioritised actions", "3,000", "Title, priority (H/M/L), affected themes, rationale, timeline. JSON output with fallback"],
        ],
        col_widths=[1.5, 1.5, 1.0, 2.5]
    )

    doc.add_heading("7.6 Document Intelligence", level=2)
    add_body(doc, "Four-stage pipeline for uploaded documents:")
    doc_pipeline = [
        ("1. Text extraction: ", "PyMuPDF for PDFs, XML parsing for DOCX files. Page-by-page extraction with newline joining."),
        ("2. Classification: ", "AI classifies into 10 categories (terms_of_reference, policy_document, financial_report, meeting_minutes, programme_specification, accreditation_report, student_survey, SOP, org_chart, other) with confidence scores. Filename-keyword fallback."),
        ("3. Structured extraction: ", "Type-specific data extraction (planned enhancement for financial data, key-value pairs, table extraction)."),
        ("4. Completeness check: ", "AI evaluates whether the document satisfies item requirements, scoring 0\u2013100 with present/missing section identification and improvement recommendations."),
    ]
    for bold, rest in doc_pipeline:
        add_bullet(doc, rest, bold_prefix=bold)

    doc.add_heading("7.7 Risk Prediction", level=2)
    add_body(doc, "Pure rule-based risk scoring engine with 6 weighted factors (no AI API calls):")
    add_styled_table(doc,
        ["Factor", "Weight", "Risk Threshold", "Description"],
        [
            ["Financial Sustainability", "25%", "Score < 40", "Low financial health score indicates funding risk"],
            ["Enrollment Trends", "20%", "Decreasing", "Declining student numbers signal partnership viability concerns"],
            ["Student Retention", "15%", "Rate < 70%", "Low retention suggests quality or support issues"],
            ["Student-Staff Ratio", "15%", "SSR > 35", "High ratios indicate understaffing"],
            ["Governance Strength", "15%", "Score < 50", "Weak governance undermines partnership quality"],
            ["Staff Qualifications", "10%", "PhD% < 20%", "Low qualification levels affect teaching quality"],
        ],
        col_widths=[1.5, 0.7, 1.0, 3.3]
    )
    doc.add_paragraph("")
    add_body(doc, "Risk levels: score \u2265 0.6 = High, 0.3\u20130.6 = Medium, < 0.3 = Low. Each contributing factor reports its raw score and weighted contribution.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 8. INFRASTRUCTURE & DEPLOYMENT
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("8. Infrastructure & Deployment", level=1)

    doc.add_heading("8.1 Docker Services", level=2)
    add_styled_table(doc,
        ["Service", "Image", "Port", "Purpose"],
        [
            ["postgres", "postgres:16-alpine", "5432", "Primary database with persistent volume"],
            ["redis", "redis:7-alpine", "6379", "Cache, Celery broker (queue 1), result backend (queue 2)"],
            ["minio", "minio/minio:latest", "9000/9001", "S3-compatible object storage for documents"],
            ["mailpit", "axllent/mailpit:latest", "1025/8025", "Development email server with web UI"],
            ["migrate", "Backend Dockerfile", "\u2014", "One-shot Alembic migration runner (blocks backend)"],
            ["backend", "Backend Dockerfile", "8000", "FastAPI application (2 Uvicorn workers)"],
            ["celery-worker", "Backend Dockerfile", "\u2014", "Celery task processor (concurrency=2)"],
            ["frontend", "Frontend Dockerfile", "3000", "Next.js standalone server"],
        ],
        col_widths=[1.2, 1.6, 0.8, 2.9]
    )

    doc.add_heading("8.2 Service Dependencies", level=2)
    dep_items = [
        "Migrate service runs alembic upgrade head before backend starts (service_completed_successfully).",
        "Backend depends on postgres (healthy), redis (healthy), minio (started), migrate (completed).",
        "Celery worker shares all backend dependencies.",
        "Frontend depends on backend (healthy). Proxies /api/* requests via Next.js rewrites (no CORS).",
        "All services use Docker health checks with 5\u201310 second intervals.",
    ]
    for d in dep_items:
        add_bullet(doc, d)

    doc.add_heading("8.3 Build Architecture", level=2)
    build_items = [
        ("Backend: ", "Multi-stage Docker build (builder \u2192 runtime). Python 3.11-slim base, non-root user, PYTHONPATH=/app."),
        ("Frontend: ", "Multi-stage Docker build (deps \u2192 builder \u2192 runner). Node 20-alpine, Next.js standalone output, BACKEND_URL resolved at build time."),
    ]
    for bold, rest in build_items:
        add_bullet(doc, rest, bold_prefix=bold)

    doc.add_heading("8.4 CI/CD Pipeline", level=2)
    add_body(doc, "GitHub Actions workflow triggered on push/PR to main with 4 parallel jobs:")
    ci_items = [
        ("Backend Lint: ", "Ruff linting and format checking."),
        ("Backend Test: ", "Pytest with PostgreSQL and Redis service containers, coverage reporting."),
        ("Frontend Lint: ", "ESLint and TypeScript type checking (--noEmit)."),
        ("Frontend Build: ", "Production build verification."),
    ]
    for bold, rest in ci_items:
        add_bullet(doc, rest, bold_prefix=bold)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 9. TESTING STRATEGY
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("9. Testing Strategy", level=1)

    doc.add_heading("9.1 Unit & Integration Tests", level=2)
    add_body(doc, "Backend tests use pytest with pytest-asyncio for async test support and factory-boy for test data generation. Coverage reporting via pytest-cov.")

    doc.add_heading("9.2 End-to-End Tests", level=2)
    add_body(doc, "11 Playwright test suites covering the complete user journey:")
    add_styled_table(doc,
        ["Suite", "Focus Area", "Primary Role"],
        [
            ["01 - Registration & Onboarding", "New user signup, email verification flow", "New user"],
            ["02 - Authentication", "Login, logout, magic links, protected routes", "All roles"],
            ["03 - Assessment Lifecycle", "Create, edit, save, submit assessment", "Tenant Admin"],
            ["04 - Assessor Journey", "Fill assessment responses, auto-save behaviour", "Assessor"],
            ["05 - Reviewer Journey", "Review, trigger scoring, generate report", "Reviewer"],
            ["06 - Platform Admin", "Global stats, tenant management", "Platform Admin"],
            ["07 - Tenant Management", "Partner CRUD, tenant settings", "Tenant Admin"],
            ["08 - User Management", "User creation, role assignment", "Tenant Admin"],
            ["09 - File Upload", "Upload policy documents, evidence files", "Assessor"],
            ["10 - Benchmarking", "View benchmarks, peer comparison charts", "Tenant Admin"],
            ["11 - Error & Edge Cases", "Cross-tenant isolation, error handling", "Various"],
        ],
        col_widths=[1.8, 2.5, 1.2]
    )

    doc.add_heading("9.3 Test Infrastructure", level=2)
    test_infra = [
        ("Pre-authenticated fixtures: ", "JWT tokens injected into localStorage, bypassing login UI for faster tests."),
        ("API seeder: ", "Direct HTTP client for setup/teardown of test data (templates, assessments, scoring)."),
        ("Mailpit integration: ", "Email token extraction for verification flow testing."),
        ("Field helpers: ", "Locators for all 12 field types with fill methods."),
        ("Wait helpers: ", "Polling utilities for auto-save and Celery job completion."),
        ("Configuration: ", "Sequential execution, 60s timeout, Chromium-only, traces and screenshots on failure."),
    ]
    for bold, rest in test_infra:
        add_bullet(doc, rest, bold_prefix=bold)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 10. NON-FUNCTIONAL REQUIREMENTS
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("10. Non-Functional Requirements", level=1)

    doc.add_heading("10.1 Security", level=2)
    sec_items = [
        "JWT HS256 tokens with 15-minute access and 7-day refresh expiry.",
        "Password hashing via bcrypt (passlib).",
        "Email verification required before login.",
        "Rate limiting on authentication endpoints (3 requests/60 seconds per email).",
        "Multi-tenant data isolation enforced at application layer on all queries.",
        "Presigned S3 URLs for file downloads (1-hour TTL, no direct storage access).",
        "CORS configuration restricted to known origins.",
        "Single-use tokens for email verification and magic links (consumed on use).",
    ]
    for s in sec_items:
        add_bullet(doc, s)

    doc.add_heading("10.2 Performance", level=2)
    perf_items = [
        "Async I/O throughout the backend (asyncpg, async SQLAlchemy, FastAPI).",
        "Celery workers for all AI operations (non-blocking, background processing).",
        "7-day in-memory cache on AI API calls (SHA256 deduplication).",
        "3-attempt exponential backoff retry on AI calls (2s\u201330s).",
        "React Query client-side caching with 60-second stale time.",
        "Auto-save debounce (2 seconds) prevents excessive API calls.",
        "2 Uvicorn workers + 2 Celery workers in production configuration.",
    ]
    for p in perf_items:
        add_bullet(doc, p)

    doc.add_heading("10.3 Scalability", level=2)
    scale_items = [
        "Stateless backend design (JWT auth, no server sessions) enables horizontal scaling.",
        "Celery workers can scale independently of API workers.",
        "PostgreSQL with JSONB supports flexible data without schema changes.",
        "S3-compatible storage (MinIO) can be replaced with AWS S3 for production scale.",
        "Redis is ephemeral and horizontally scalable.",
    ]
    for s in scale_items:
        add_bullet(doc, s)

    doc.add_heading("10.4 Reliability", level=2)
    rel_items = [
        "Docker health checks prevent cascading failures between services.",
        "Migration service blocks backend startup until schema is current.",
        "AI pipeline graceful degradation: filename-keyword fallback for document classification, rule-based consistency checks when AI unavailable.",
        "Celery task_acks_late ensures tasks are re-queued if worker crashes mid-execution.",
        "Single-prefetch (worker_prefetch_multiplier=1) prevents task loss.",
    ]
    for r in rel_items:
        add_bullet(doc, r)

    doc.add_heading("10.5 Observability", level=2)
    obs_items = [
        "Structured logging via structlog throughout the backend.",
        "AI API cost tracking per call (input/output token counts and estimated USD cost).",
        "AI Job model provides full audit trail (queued \u2192 processing \u2192 completed/failed with timestamps and result data).",
        "Celery task_track_started enables real-time task status monitoring.",
    ]
    for o in obs_items:
        add_bullet(doc, o)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 11. ASSESSMENT TEMPLATE
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("11. Assessment Template Specification", level=1)
    add_body(doc, "The TNE Quality Assessment v1.0 template contains 52 items across 5 weighted themes:")

    add_styled_table(doc,
        ["Theme", "Code", "Weight", "Items", "Focus Areas"],
        [
            ["Teaching & Learning", "TL", "25%", "15", "Programmes, staff qualifications, teaching methods, student-staff ratios, flying faculty"],
            ["Student Experience & Outcomes", "SE", "25%", "12", "Enrollment, retention, graduation, employment, student support, satisfaction"],
            ["Governance & Quality Assurance", "GV", "20%", "10", "QA frameworks, accreditation, policies, academic standards, governance structures"],
            ["Impact & Engagement", "IM", "15%", "8", "Research output, community engagement, industry partnerships, innovation"],
            ["Financial Sustainability", "FN", "15%", "7", "Revenue, expenditure, financial planning, salary benchmarking, sustainability"],
        ],
        col_widths=[1.8, 0.5, 0.7, 0.5, 3.0]
    )

    doc.add_paragraph("")
    add_body(doc,
        "Each item includes a unique code (e.g., TL01, SE05, GV03), a human-readable label, "
        "optional description and help text, field type and configuration, scoring rubric, weight, "
        "and display ordering. Items may have conditional dependencies (depends_on) and validation rules."
    )

    doc.add_heading("11.1 Auto-Calculated Fields", level=2)
    add_styled_table(doc,
        ["Field Code", "Calculation", "Dependencies"],
        [
            ["TL_SSR", "Student-Staff Ratio = total_students / total_academic_staff", "TL03, TL06"],
            ["TL_PHD_PCT", "PhD% = phd_staff / total_academic_staff \u00d7 100", "TL07, TL06"],
            ["TL_FLYING_PCT", "Flying Faculty% = flying_faculty / total_academic_staff \u00d7 100", "TL09, TL06"],
            ["SE_RETENTION", "Retention Rate = completed / enrolled \u00d7 100", "SE01, SE02"],
            ["IM_EMPLOYMENT", "Employment Rate = employed / graduates \u00d7 100", "IM01, IM02"],
        ],
        col_widths=[1.3, 3.2, 2.0]
    )

    doc.add_heading("11.2 Scoring Formula", level=2)
    add_body(doc, "The overall assessment score is calculated as follows:")
    score_steps = [
        "Each item is scored 0\u2013100 by its type-specific scorer.",
        "Theme normalised score = weighted average of item scores within the theme: \u03a3(item_score \u00d7 item_weight) / \u03a3(item_weight).",
        "Theme weighted score = normalised_score \u00d7 theme_weight (TL: 0.25, SE: 0.25, GV: 0.20, IM: 0.15, FN: 0.15).",
        "Overall score = \u03a3(all theme weighted scores), yielding a 0\u2013100 final score.",
    ]
    for s in score_steps:
        add_bullet(doc, s)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 12. GLOSSARY
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("12. Glossary", level=1)
    add_styled_table(doc,
        ["Term", "Definition"],
        [
            ["TNE", "Transnational Education \u2014 educational programmes delivered in a country other than the awarding institution's home country"],
            ["Tenant", "An organisation (institution) using the platform; each tenant has isolated data and users"],
            ["Assessment", "A structured self-evaluation completed by an institution for a specific academic year"],
            ["Theme", "One of 5 major assessment categories (Teaching & Learning, Student Experience, Governance, Impact, Financial)"],
            ["Item", "An individual question or data field within a theme (52 total across all themes)"],
            ["Rubric", "Scoring criteria used to evaluate a response, either AI-based (4 dimensions) or rule-based (ranges)"],
            ["SSR", "Student-Staff Ratio \u2014 total enrolled students divided by total academic staff"],
            ["Flying Faculty", "Academic staff who travel from the home institution to deliver teaching at partner sites"],
            ["Benchmark", "Statistical comparison of an institution's scores against peer institutions"],
            ["Percentile", "The percentage of institutions scoring at or below a given value (P50 = median)"],
            ["Celery", "Python distributed task queue used for asynchronous background processing"],
            ["JWT", "JSON Web Token \u2014 stateless authentication mechanism used for API access"],
            ["MinIO", "S3-compatible open-source object storage system"],
            ["JSONB", "PostgreSQL binary JSON column type supporting indexing and querying"],
        ],
        col_widths=[1.3, 5.2]
    )

    # ── Footer ──
    doc.add_paragraph("")
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("\u2014 End of Document \u2014")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.italic = True

    return doc


if __name__ == "__main__":
    output_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "docs",
        "TNE_Platform_Scoping_Requirements.docx"
    )
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = build_document()
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
